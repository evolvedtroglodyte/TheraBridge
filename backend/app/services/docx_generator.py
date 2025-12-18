"""DOCX generation service using python-docx"""
import logging
from io import BytesIO
from typing import Dict, Any, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DOCXGeneratorService:
    """Service for generating DOCX documents"""

    def __init__(self):
        """Initialize DOCX generator"""
        pass

    async def generate_progress_report(
        self,
        patient: Dict[str, Any],
        therapist: Dict[str, Any],
        goals: List[Dict[str, Any]],
        sessions: List[Dict[str, Any]],
        date_range: Dict[str, datetime],
        include_sections: Dict[str, bool]
    ) -> bytes:
        """
        Generate progress report DOCX.

        Args:
            patient: Patient data
            therapist: Therapist data
            goals: Treatment goals with progress
            sessions: Session summaries
            date_range: {start_date, end_date}
            include_sections: Which sections to include

        Returns:
            DOCX file as bytes
        """
        try:
            logger.info("Starting DOCX progress report generation")

            doc = Document()

            # Add confidentiality notice
            heading = doc.add_paragraph("CONFIDENTIAL - PROTECTED HEALTH INFORMATION")
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.color.rgb = RGBColor(211, 47, 47)
            heading.runs[0].font.bold = True

            # Add title
            title = doc.add_heading('Progress Report', 0)

            # Add header information
            doc.add_paragraph(f"Patient: {patient['full_name']}")
            doc.add_paragraph(
                f"Period: {date_range['start_date'].strftime('%B %d, %Y')} - "
                f"{date_range['end_date'].strftime('%B %d, %Y')}"
            )
            doc.add_paragraph(f"Prepared by: {therapist['full_name']}")
            doc.add_paragraph("")  # Blank line

            # Patient information section
            if include_sections.get('patient_info', True):
                doc.add_heading('Patient Information', 1)
                doc.add_paragraph(f"Name: {patient['full_name']}")
                doc.add_paragraph(f"Email: {patient.get('email', 'N/A')}")
                doc.add_paragraph("")

            # Treatment goals section
            if include_sections.get('treatment_goals', True) and goals:
                doc.add_heading('Treatment Goals & Progress', 1)

                # Create table
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'

                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Goal'
                hdr_cells[1].text = 'Baseline'
                hdr_cells[2].text = 'Current'
                hdr_cells[3].text = 'Progress'

                # Data rows
                for goal in goals:
                    row_cells = table.add_row().cells
                    row_cells[0].text = goal['description']
                    row_cells[1].text = str(goal.get('baseline', 'N/A'))
                    row_cells[2].text = str(goal.get('current', 'N/A'))
                    row_cells[3].text = f"{goal.get('progress_percentage', 0)}%"

                doc.add_paragraph("")

            # Session summary section
            if include_sections.get('session_summary', True):
                doc.add_heading('Session Summary', 1)
                doc.add_paragraph(f"Total Sessions: {len(sessions)}")

                if sessions:
                    avg_duration = sum(s.get('duration_minutes', 0) for s in sessions) / len(sessions)
                    doc.add_paragraph(f"Average Duration: {avg_duration:.1f} minutes")

                doc.add_paragraph("")

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.read()

            logger.info(
                f"DOCX generated successfully",
                extra={"size_kb": len(docx_bytes) / 1024}
            )

            return docx_bytes

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}", exc_info=True)
            raise ValueError(f"Failed to generate DOCX: {str(e)}")

    async def generate_treatment_summary(
        self,
        patient: Dict[str, Any],
        therapist: Dict[str, Any],
        treatment_data: Dict[str, Any],
        include_sections: Dict[str, bool]
    ) -> bytes:
        """Generate treatment summary DOCX"""
        # Similar structure to progress_report
        # ... implementation ...
        pass

    async def generate_timeline_export(
        self,
        patient: Dict[str, Any],
        therapist: Dict[str, Any],
        events: List[Dict[str, Any]],
        summary: Dict[str, Any],
        date_range: Dict[str, Any]
    ) -> bytes:
        """
        Generate patient timeline DOCX export.

        Args:
            patient: Patient data
            therapist: Therapist data
            events: Timeline events list
            summary: Timeline summary data
            date_range: {start_date, end_date} or None values

        Returns:
            DOCX file as bytes
        """
        try:
            logger.info("Starting DOCX timeline export generation")

            doc = Document()

            # Add confidentiality notice
            heading = doc.add_paragraph("CONFIDENTIAL - PROTECTED HEALTH INFORMATION")
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.color.rgb = RGBColor(211, 47, 47)
            heading.runs[0].font.bold = True

            # Add title
            title = doc.add_heading('Patient Timeline Export', 0)

            # Add header information
            doc.add_paragraph(f"Patient: {patient['full_name']}")

            if date_range.get('start_date') and date_range.get('end_date'):
                doc.add_paragraph(
                    f"Period: {date_range['start_date'].strftime('%B %d, %Y')} - "
                    f"{date_range['end_date'].strftime('%B %d, %Y')}"
                )
            else:
                doc.add_paragraph("Period: All Time")

            doc.add_paragraph(f"Prepared by: {therapist['full_name']}")
            doc.add_paragraph(f"Export Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("")  # Blank line

            # Summary section
            doc.add_heading('Timeline Summary', 1)
            doc.add_paragraph(f"Total Events: {len(events)}")
            doc.add_paragraph(f"Total Sessions: {summary.get('total_sessions', 0)}")
            doc.add_paragraph(f"Milestones Achieved: {summary.get('milestones_achieved', 0)}")

            if summary.get('first_session'):
                doc.add_paragraph(f"First Session: {summary['first_session']}")
            if summary.get('last_session'):
                doc.add_paragraph(f"Last Session: {summary['last_session']}")

            doc.add_paragraph("")

            # Events by type breakdown
            if summary.get('events_by_type'):
                doc.add_heading('Events by Type', 2)
                for event_type, count in summary['events_by_type'].items():
                    doc.add_paragraph(f"{event_type.capitalize()}: {count}", style='List Bullet')
                doc.add_paragraph("")

            # Timeline events section
            doc.add_heading('Timeline Events', 1)

            if not events:
                doc.add_paragraph("No events found for the specified criteria.")
            else:
                # Group events by importance for better organization
                milestone_events = [e for e in events if e['importance'] == 'milestone']
                high_events = [e for e in events if e['importance'] == 'high']
                normal_events = [e for e in events if e['importance'] == 'normal']
                low_events = [e for e in events if e['importance'] == 'low']

                # Display milestones first
                if milestone_events:
                    doc.add_heading('Milestones', 2)
                    for event in milestone_events:
                        self._add_event_to_doc(doc, event)
                    doc.add_paragraph("")

                # Display high importance events
                if high_events:
                    doc.add_heading('High Importance Events', 2)
                    for event in high_events:
                        self._add_event_to_doc(doc, event)
                    doc.add_paragraph("")

                # Display normal events
                if normal_events:
                    doc.add_heading('Standard Events', 2)
                    for event in normal_events:
                        self._add_event_to_doc(doc, event)
                    doc.add_paragraph("")

                # Display low importance events
                if low_events:
                    doc.add_heading('Additional Events', 2)
                    for event in low_events:
                        self._add_event_to_doc(doc, event)

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.read()

            logger.info(
                f"Timeline DOCX generated successfully",
                extra={"size_kb": len(docx_bytes) / 1024, "event_count": len(events)}
            )

            return docx_bytes

        except Exception as e:
            logger.error(f"Timeline DOCX generation failed: {e}", exc_info=True)
            raise ValueError(f"Failed to generate timeline DOCX: {str(e)}")

    def _add_event_to_doc(self, doc: Document, event: Dict[str, Any]) -> None:
        """
        Add a single event to the document.

        Args:
            doc: Document instance
            event: Event data dictionary
        """
        # Event title with date
        event_date = event['event_date']
        if isinstance(event_date, str):
            date_str = event_date
        else:
            date_str = event_date.strftime('%B %d, %Y')

        event_para = doc.add_paragraph()
        event_para.add_run(f"{event['title']}").bold = True
        event_para.add_run(f" - {date_str}")

        # Event type badge
        type_para = doc.add_paragraph(f"Type: {event['event_type'].capitalize()}", style='List Bullet')

        # Description if available
        if event.get('description'):
            doc.add_paragraph(event['description'])

        # Metadata if available
        if event.get('metadata'):
            metadata = event['metadata']
            if isinstance(metadata, dict):
                if metadata.get('topics'):
                    topics_str = ", ".join(metadata['topics'])
                    doc.add_paragraph(f"Topics: {topics_str}", style='List Bullet 2')
                if metadata.get('mood'):
                    doc.add_paragraph(f"Mood: {metadata['mood']}", style='List Bullet 2')

        doc.add_paragraph("")  # Blank line after event


def get_docx_generator() -> DOCXGeneratorService:
    """FastAPI dependency to provide DOCX generator service"""
    return DOCXGeneratorService()
